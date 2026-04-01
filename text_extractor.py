#!/usr/bin/env python3
"""
Extract text from PDF, DOCX, ODT, and TXT files in a directory tree.

Usage:
    text_extractor.py <directory> [--output=<file>] [--verbose] [--include=<extensions>]
    text_extractor.py -h | --help

Options:
    -o, --output=<file>       Output JSON file [default: extracted_text.json]
    -v, --verbose             Print progress information
    --include=<extensions>    Comma-separated list of extensions to include 
                             [default: pdf,docx,odt,txt]
    -h, --help               Show this help message
"""

import os
import sys
import json
from pathlib import Path
from docopt_ng import docopt

try:
    import PyPDF2
except ImportError:
    print("Error: PyPDF2 is required. Install it with: pip install PyPDF2")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("Error: python-docx is required for .docx files. Install with: pip install python-docx")
    sys.exit(1)

try:
    from odf import text, teletype
    from odf.opendocument import load
except ImportError:
    print("Error: odfpy is required for .odt files. Install with: pip install odfpy")
    sys.exit(1)


def extract_pdf_text(pdf_path):
    """Extract text from a single PDF file."""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            
            return '\n'.join(text) if text else ""
    
    except Exception as e:
        return f"Error extracting text: {str(e)}"


def extract_docx_text(docx_path):
    """Extract text from a DOCX file."""
    try:
        doc = Document(docx_path)
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text.append(paragraph.text)
        
        return '\n'.join(text) if text else ""
    
    except Exception as e:
        return f"Error extracting text: {str(e)}"


def extract_odt_text(odt_path):
    """Extract text from an ODT file."""
    try:
        doc = load(odt_path)
        texts = doc.getElementsByType(text.P)
        
        extracted_text = []
        for element in texts:
            extracted_text.append(teletype.extractText(element))
        
        return '\n'.join(extracted_text) if extracted_text else ""
    
    except Exception as e:
        return f"Error extracting text: {str(e)}"


def extract_txt_text(txt_path):
    """Extract text from a TXT file."""
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try different encoding if UTF-8 fails
        try:
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    except Exception as e:
        return f"Error extracting text: {str(e)}"


def find_files_by_extensions(directory, extensions):
    """Recursively find all files with specified extensions in directory."""
    files = []
    for root, dirs, files_in_dir in os.walk(directory):
        for file in files_in_dir:
            file_ext = Path(file).suffix.lower().lstrip('.')
            if file_ext in extensions:
                files.append(Path(root) / file)
    return files


def get_extractor(extension):
    """Return the appropriate extractor function for the file extension."""
    extractors = {
        'pdf': extract_pdf_text,
        'docx': extract_docx_text,
        'odt': extract_odt_text,
        'txt': extract_txt_text,
    }
    return extractors.get(extension.lower())


def main():
    args = docopt(__doc__)
    
    directory = args['<directory>']
    output_file = args['--output']
    verbose = args['--verbose']
    include_extensions = [ext.strip().lower() for ext in args['--include'].split(',')]
    
    # Validate directory
    if not os.path.isdir(directory):
        print(f"Error: Directory '{directory}' does not exist.")
        sys.exit(1)
    
    if verbose:
        print(f"Searching for files in: {directory}")
        print(f"Including extensions: {', '.join(include_extensions)}")
    
    # Find all files with specified extensions
    files = find_files_by_extensions(directory, include_extensions)
    
    if not files:
        print(f"No files found with extensions: {', '.join(include_extensions)}")
        sys.exit(0)
    
    if verbose:
        print(f"Found {len(files)} files to process")
    
    # Extract text from each file
    extracted_data = []
    supported_extensions = set(['pdf', 'docx', 'odt', 'txt'])
    
    for i, file_path in enumerate(files, 1):
        if verbose:
            print(f"[{i}/{len(files)}] Processing: {file_path}")
        
        extension = file_path.suffix.lower().lstrip('.')
        
        # Check if extension is supported
        if extension not in supported_extensions:
            text = f"Error: Unsupported file type '{extension}'"
        else:
            extractor = get_extractor(extension)
            text = extractor(file_path)
        
        extracted_data.append({
            'filename': str(file_path),
            'extension': extension,
            'text': text
        })
    
    # Save to JSON file
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(extracted_data, f, indent=2, ensure_ascii=False)
    
    if verbose:
        print(f"\nExtraction complete. Results saved to: {output_file}")
        print(f"Total files processed: {len(extracted_data)}")
        
        # Show statistics
        ext_counts = {}
        for item in extracted_data:
            ext = item['extension']
            ext_counts[ext] = ext_counts.get(ext, 0) + 1
        
        print("\nFiles processed by type:")
        for ext, count in ext_counts.items():
            print(f"  .{ext}: {count} files")


if __name__ == '__main__':
    main()